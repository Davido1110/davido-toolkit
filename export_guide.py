from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    # light grey background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F2937')
    pPr.append(shd)
    run.font.color.rgb = RGBColor(0x4a, 0xde, 0x80)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    # data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            if ci == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
                run.font.name = 'Courier New'
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
t = title.add_run('PO Calculator — How it works / Hướng dẫn')
t.bold = True
t.font.size = Pt(18)
t.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
title.paragraph_format.space_after = Pt(4)

sub = doc.add_paragraph('Tài liệu mô tả logic tính toán và cách sử dụng công cụ PO Calculator.')
sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
sub.paragraph_format.space_after = Pt(16)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. ABBREVIATIONS
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Giải thích viết tắt / Abbreviations')

abbrev_rows = [
    ('AVG',           'Doanh số bán trung bình / tuần',  'Average Weekly Sales',       'Trung bình số lượng bán ra mỗi tuần trong 12 tuần gần nhất. / Pre-calculated 12-week average of units sold per week.'),
    ('SOH',           'Tồn kho hiện tại',                'Stock On Hand',              'Số lượng hàng hiện đang có trong kho. / Current quantity physically available in the warehouse.'),
    ('OO',            'Đang trên đường về',              'On Order',                   'Số lượng đã đặt hàng, đang sản xuất hoặc vận chuyển, chưa về kho. / Ordered from supplier, in production or transit, not yet received.'),
    ('LT',            'Thời gian giao hàng',             'Lead Time',                  'Số tuần từ lúc đặt hàng đến khi nhận được hàng tại kho. / Weeks from placing an order to receiving it.'),
    ('D',             'Hệ số nhu cầu',                   'Demand Factor',              'Số tuần hàng muốn đặt mỗi lần đặt hàng. / Weeks of supply to order per purchase order.'),
    ('QTY',           'Số lượng cần đặt',                'Order Quantity',             'Số lượng đề xuất cần đặt mua. / Suggested number of units to order now.'),
    ('Projected SOH', 'Tồn kho dự kiến',                 'Projected Stock On Hand',    'Ước tính tồn kho tại tuần thứ LT (khi đơn hàng mới về). / Estimated warehouse stock at week LT when new order arrives.'),
    ('SKU',           'Mã hàng hóa',                     'Stock Keeping Unit',         'Mã định danh duy nhất cho từng sản phẩm / biến thể. / Unique product identifier per distinct product/variant.'),
]

add_table(
    ['Viết tắt / Term', 'Tên tiếng Việt', 'Full Name (EN)', 'Mô tả / Description'],
    abbrev_rows,
    col_widths=[1.1, 1.4, 1.5, 2.8],
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. FORMULAS
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Logic tính toán / Calculation Logic')

h2('Số lượng cần đặt (QTY) / Order Quantity')
code_block('remaining = max(0, SOH + OO − AVG × LT)')
code_block('QTY       = max(0, AVG × D − remaining)')

explain_qty = [
    ('AVG × D',                     '— Số lượng mục tiêu cần đặt: D tuần doanh số. / Target order size: D weeks of supply.'),
    ('max(0, SOH + OO − AVG × LT)', '— Tồn kho còn lại khi hàng về (0 nếu sẽ hết hàng trong thời gian chờ). / Stock remaining at arrival, floored at 0 if stockout during lead time.'),
    ('max(0, …)',                   '— Không cho kết quả âm; nếu tồn kho dư thì không cần đặt. / Floors at zero; no order needed if remaining stock already covers D weeks.'),
]
for term, desc in explain_qty:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(term)
    r1.bold = True
    r1.font.name = 'Courier New'
    r1.font.size = Pt(10)
    p.add_run(f'  {desc}').font.size = Pt(10)

doc.add_paragraph()
h2('Tồn kho dự kiến (Projected SOH)')
code_block('Projected SOH = max(0, SOH + OO − AVG × LT) + QTY')

explain_proj = [
    ('max(0, SOH + OO − AVG × LT)', '— Tồn kho thực tế còn lại khi hàng về (không thể âm). / Stock physically remaining at arrival, floored at 0.'),
    ('+ QTY',                       '— Cộng đơn hàng mới vừa về kho. / Plus the new order arriving at the warehouse.'),
]
for term, desc in explain_proj:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(term)
    r1.bold = True
    r1.font.name = 'Courier New'
    r1.font.size = Pt(10)
    p.add_run(f'  {desc}').font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. WORKED EXAMPLE
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Ví dụ minh họa / Worked Example')

h2('Dữ liệu đầu vào / Inputs')
add_table(
    ['Tham số', 'Giá trị'],
    [('AVG', '105'), ('SOH', '200'), ('OO', '150'), ('LT', '12 tuần / weeks'), ('D', '12 tuần / weeks')],
    col_widths=[1.5, 4.3],
)

h2('Tính từng bước / Step-by-step')
steps = [
    '// Tồn kho còn lại khi hàng về: max(0, 200 + 150 − 105 × 12) = max(0, 350 − 1260) = 0',
    'QTY = max(0, 105 × 12 − 0) = max(0, 1260) = 1260',
    '',
    'Projected SOH = max(0, 200 + 150 − 105 × 12) + 1260',
    '              = 0 + 1260',
    '              = 1260   ✓  (= D weeks of stock when order arrives)',
]
for s in steps:
    p = doc.add_paragraph()
    r = p.add_run(s)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)
    if s.startswith('//'):
        r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    elif '✓' in s:
        r.bold = True
        r.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(1)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. ROW COLORS
# ══════════════════════════════════════════════════════════════════════════════
h1('4. Màu sắc các dòng / Row Color Coding')

colors_data = [
    ('Đỏ / Red',     'SOH + OO < AVG × LT',  'Sẽ hết hàng trong thời gian chờ đơn hàng mới. / Stock will run out before the new order arrives.'),
    ('Vàng / Yellow','QTY > 0',              'Tồn kho đủ trong thời gian chờ nhưng cần đặt thêm. / Stock lasts through lead time but a purchase order is needed.'),
    ('Trắng / White','QTY = 0',              'Tồn kho hiện tại đủ phủ toàn bộ. / Current stock + on-order is sufficient.'),
]
add_table(
    ['Màu / Color', 'Điều kiện / Condition', 'Ý nghĩa / Meaning'],
    colors_data,
    col_widths=[1.2, 1.8, 4.3],
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. INPUT FILE FORMAT
# ══════════════════════════════════════════════════════════════════════════════
h1('5. Định dạng file đầu vào / Input File Format')
body('Tải lên file .xlsx hoặc .xls. Dòng 1 là tiêu đề và bị bỏ qua. Dữ liệu bắt đầu từ dòng 2.')
body('Upload an .xlsx or .xls file. Row 1 is treated as a header and skipped. Data begins from row 2.')
doc.add_paragraph()

add_table(
    ['Cột / Col', 'Trường / Field', 'Kiểu / Type', 'Ví dụ / Example'],
    [
        ('A', 'Mã hàng / SKU Code',     'Văn bản / Text',   'SKU-001'),
        ('B', 'Nhóm hàng / Category',   'Văn bản / Text',   'Electronics > Phones > iPhone'),
        ('C', 'Tên hàng / Product Name','Văn bản / Text',   'iPhone 15 Pro'),
        ('D', 'Thuộc tính / Attribute', 'Văn bản / Text',   '128GB Black'),
        ('E', 'AVG',                    'Số / Number',       '105'),
        ('F', 'SOH',                    'Số / Number',       '200'),
        ('G', 'OO',                     'Số / Number',       '150'),
    ],
    col_widths=[0.5, 2.0, 1.2, 3.1],
)

note = doc.add_paragraph('* Nhóm hàng dùng dấu > để phân cấp (ví dụ: L1 > L2 > L3). Hỗ trợ tối đa 3 cấp.')
note.runs[0].font.size = Pt(9)
note.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

# ── Save ─────────────────────────────────────────────────────────────────────
out = '/Users/Davido/Documents/davido-toolkit/PO_Calculator_Guide.docx'
doc.save(out)
print(f'Saved: {out}')
